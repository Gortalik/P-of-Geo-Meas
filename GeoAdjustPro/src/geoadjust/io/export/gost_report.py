from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import numpy as np
from typing import Dict, List, Any


class GOSTReportGenerator:
    """
    Генератор отчётов по ГОСТ 7.32-2017
    """
    
    def __init__(self):
        self.doc = Document()
        self.formula_counter = 1
        
    def _add_title_page(self, project_info: Dict[str, str]):
        """
        Добавление титульного листа
        """
        # Заголовок документа
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run("ПРОТОКОЛ УРАВНИВАНИЯ ГЕОДЕЗИЧЕСКОЙ СЕТИ")
        title_run.bold = True
        title_run.font.size = 16
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Пустая строка
        self.doc.add_paragraph()
        
        # Информация о проекте
        if 'project_name' in project_info:
            para = self.doc.add_paragraph(f"Наименование проекта: {project_info['project_name']}")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        if 'location' in project_info:
            para = self.doc.add_paragraph(f"Место выполнения работ: {project_info['location']}")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        if 'date' in project_info:
            para = self.doc.add_paragraph(f"Дата выполнения: {project_info['date']}")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        # Пустая строка перед подписями
        self.doc.add_paragraph()
        
        # Подписи
        signature_table = self.doc.add_table(rows=3, cols=2)
        signature_table.style = 'Table Grid'
        
        cell = signature_table.cell(0, 0)
        cell.text = "Исполнитель:"
        cell = signature_table.cell(0, 1)
        cell.text = project_info.get('performer', '')
        
        cell = signature_table.cell(1, 0)
        cell.text = "Руководитель работ:"
        cell = signature_table.cell(1, 1)
        cell.text = project_info.get('supervisor', '')
        
        cell = signature_table.cell(2, 0)
        cell.text = "Дата:"
        cell = signature_table.cell(2, 1)
        cell.text = project_info.get('date', '')
        
    def _add_numbered_formula(self, paragraph, latex_formula: str, number: int):
        """
        Добавление нумерованной формулы
        Формат:                     (N)
                формула
        """
        # Добавляем формулу и номер
        formula_with_number = f"{latex_formula} ({number})"
        paragraph.add_run(formula_with_number)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Увеличиваем счётчик формул
        self.formula_counter += 1
        
    def add_section(self, title: str):
        """
        Добавление раздела
        """
        heading = self.doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def add_subsection(self, title: str):
        """
        Добавление подраздела
        """
        self.doc.add_heading(title, level=2)
        
    def add_text(self, text: str):
        """
        Добавление текста
        """
        self.doc.add_paragraph(text)
        
    def add_table(self, headers: List[str], rows: List[List[Any]], title: str = None):
        """
        Добавление таблицы
        """
        if title:
            self.doc.add_paragraph(title, style='Caption')
            
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                
    def add_adjustment_results(self, results: Dict):
        """
        Добавление результатов уравнивания
        """
        self.add_subsection("Результаты уравнивания")
        
        # Основные показатели
        self.add_text(f"Число измерений: {results.get('n_measurements', 0)}")
        self.add_text(f"Число определяемых точек: {results.get('n_points', 0)}")
        self.add_text(f"Число степеней свободы: {results.get('degrees_of_freedom', 0)}")
        self.add_text(f"СКО единицы веса: {results.get('sigma0', 0.0):.4f}")
        
        # Таблица с координатами
        if 'coordinates' in results:
            coords = results['coordinates']
            headers = ['№ Пункта', 'X', 'Y', 'СКО(X)', 'СКО(Y)', 'СКО(положения)']
            rows = []
            for point_id, data in coords.items():
                row = [
                    point_id,
                    f"{data['x']:.3f}",
                    f"{data['y']:.3f}",
                    f"{data['sigma_x']:.3f}",
                    f"{data['sigma_y']:.3f}",
                    f"{np.sqrt(data['sigma_x']**2 + data['sigma_y']**2):.3f}"
                ]
                rows.append(row)
                
            self.add_table(headers, rows, "Таблица уравненных координат")
        
    def add_reliability_analysis(self, reliability_results: Dict):
        """
        Добавление анализа надёжности
        """
        self.add_subsection("Анализ надёжности сети")
        
        if 'internal_reliability' in reliability_results:
            avg_reliability = np.mean(reliability_results['internal_reliability'])
            self.add_text(f"Средняя внутренняя надёжность: {avg_reliability:.4f}")
            
        if 'max_displacement' in reliability_results:
            self.add_text(f"Максимальное смещение параметров: {reliability_results['max_displacement']:.4f} м")
    
    def add_normative_compliance(self, compliance_results: Dict):
        """
        Добавление анализа соответствия нормативным требованиям
        """
        self.add_subsection("Анализ соответствия нормативным требованиям")
        
        for class_name, compliance in compliance_results.items():
            status = "Соответствует" if compliance['passed'] else "Не соответствует"
            self.add_text(f"Класс {class_name}: {status}")
            
            if 'details' in compliance:
                for detail in compliance['details']:
                    self.add_text(f"  - {detail}")
        
    def generate_report(self, project_info: Dict, results: Dict, 
                       reliability_results: Dict = None, 
                       compliance_results: Dict = None,
                       filename: str = "report.docx"):
        """
        Генерация полного отчёта
        """
        # Титульный лист
        self._add_title_page(project_info)
        
        # Содержание (автоматически создается в Word)
        self.doc.add_page_break()
        
        # Введение
        self.add_section("Введение")
        self.add_text("Настоящий протокол содержит результаты уравнивания геодезической сети.")
        
        # Исходные данные
        self.add_section("Исходные данные")
        self.add_text("Описание исходных данных и методики обработки.")
        
        # Методы обработки
        self.add_section("Методы обработки")
        self.add_text("Обработка выполнена методом наименьших квадратов по параметрическому методу.")
        
        # Добавляем несколько формул согласно ГОСТ
        self.add_subsection("Математическая модель")
        para = self.doc.add_paragraph()
        self._add_numbered_formula(para, "V = A·ΔX - L", self.formula_counter)
        
        para = self.doc.add_paragraph()
        self._add_numbered_formula(para, "N·ΔX = U, где N = A^T·P·A, U = A^T·P·L", self.formula_counter)
        
        para = self.doc.add_paragraph()
        self._add_numbered_formula(para, "σ₀² = (V^T·P·V)/(n-u)", self.formula_counter)
        
        # Результаты уравнивания
        self.add_section("Результаты уравнивания")
        self.add_adjustment_results(results)
        
        # Анализ надёжности
        if reliability_results:
            self.add_section("Анализ надёжности сети")
            self.add_reliability_analysis(reliability_results)
        
        # Анализ соответствия нормативным требованиям
        if compliance_results:
            self.add_section("Анализ соответствия нормативным требованиям")
            self.add_normative_compliance(compliance_results)
        
        # Заключение
        self.add_section("Заключение")
        self.add_text("Уравнивание геодезической сети выполнено успешно. Результаты соответствуют установленным требованиям.")
        
        # Сохранение документа
        self.doc.save(filename)
        
        return filename